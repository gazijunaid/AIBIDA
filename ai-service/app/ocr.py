"""
Module 3 - OCR & Document Processing Engine
Extracts raw text from PDFs (native + scanned), images, DOCX, XLSX and CSV files.
"""
import os
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
from pypdf import PdfReader
import docx
import openpyxl
import pandas as pd

from app.config import TESSERACT_CMD

pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD


def _ocr_image(image: Image.Image) -> str:
    # lang="eng" by default; pass lang="eng+hin+urd" etc. for multi-language OCR (bonus feature)
    return pytesseract.image_to_string(image)


def extract_from_image(file_path: str) -> str:
    image = Image.open(file_path)
    return _ocr_image(image)


def extract_from_pdf(file_path: str) -> str:
    """Try native text extraction first; fall back to OCR for scanned PDFs."""
    text_parts = []
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    except Exception:
        pass

    combined = "\n".join(text_parts).strip()

    # If native extraction yielded little/no text, the PDF is likely scanned -> run OCR
    if len(combined) < 30:
        try:
            pages = convert_from_path(file_path, dpi=300)
            ocr_text = "\n".join(_ocr_image(p) for p in pages)
            return ocr_text
        except Exception as e:
            return combined + f"\n[OCR fallback failed: {e}]"

    return combined


def extract_from_docx(file_path: str) -> str:
    document = docx.Document(file_path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_from_xlsx(file_path: str) -> str:
    wb = openpyxl.load_workbook(file_path, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"--- Sheet: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            row_vals = [str(c) if c is not None else "" for c in row]
            if any(row_vals):
                parts.append(" | ".join(row_vals))
    return "\n".join(parts)


def extract_from_csv(file_path: str) -> str:
    df = pd.read_csv(file_path)
    return df.to_csv(index=False)


def extract_text(file_path: str, file_type: str) -> str:
    """Dispatch to the correct extractor based on detected file_type."""
    try:
        if file_type == "image":
            return extract_from_image(file_path)
        if file_type == "pdf":
            return extract_from_pdf(file_path)
        if file_type == "docx":
            return extract_from_docx(file_path)
        if file_type == "xlsx":
            return extract_from_xlsx(file_path)
        if file_type == "csv":
            return extract_from_csv(file_path)
        # Fallback: try reading as plain text
        with open(file_path, "r", errors="ignore") as f:
            return f.read()
    except Exception as e:
        return f"[Extraction error: {e}]"
